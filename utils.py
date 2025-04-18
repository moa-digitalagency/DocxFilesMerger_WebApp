import os
import shutil
import zipfile
import json
import time
import threading
import traceback
from datetime import datetime, timedelta
from pathlib import Path
import subprocess
import sys
import tempfile

# Import des bibliothèques de traitement de documents
try:
    import docx
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Bibliothèque python-docx non installée. Certaines fonctionnalités peuvent ne pas fonctionner correctement.")

def save_status(status_dir, status_data):
    """Save processing status to a JSON file"""
    if not status_dir:
        return

    os.makedirs(status_dir, exist_ok=True)
    status_file = os.path.join(status_dir, 'status.json')
    
    try:
        with open(status_file, 'w') as f:
            json.dump(status_data, f)
    except Exception as e:
        print(f"Erreur lors de la sauvegarde du statut: {str(e)}")

def extract_doc_files(zip_path, extract_dir):
    """Extract all .doc and .docx files from a zip file"""
    # Créer le dossier d'extraction s'il n'existe pas
    os.makedirs(extract_dir, exist_ok=True)
    
    # Liste pour stocker les chemins vers les fichiers extraits
    extracted_files = []
    
    # Ouvrir le fichier ZIP
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        # Parcourir tous les fichiers dans le ZIP
        for file_info in zip_ref.infolist():
            # Ignorer les dossiers
            if file_info.filename.endswith('/'):
                continue
            
            # Vérifier si le fichier est un .doc ou .docx
            if file_info.filename.lower().endswith(('.doc', '.docx')):
                # Extraire uniquement le nom du fichier sans les dossiers
                filename = os.path.basename(file_info.filename)
                # Créer un chemin de destination
                dest_path = os.path.join(extract_dir, filename)
                
                # Extraire le fichier
                with zip_ref.open(file_info) as source, open(dest_path, 'wb') as dest:
                    shutil.copyfileobj(source, dest)
                
                # Ajouter le chemin à la liste des fichiers extraits
                extracted_files.append(dest_path)
    
    return extracted_files

def convert_doc_to_docx(doc_path, output_dir):
    """
    Convert a .doc file to .docx format
    
    This function tries multiple methods to convert .doc to .docx:
    1. LibreOffice conversion (if available) - Meilleure option pour préserver la mise en forme
    2. python-docx direct loading (works for some .docx disguised as .doc)
    3. Antiword/catdoc for extracting text content (for old .doc files)
    4. Creating a placeholder document if all else fails
    """
    import subprocess
    import os
    from docx import Document
    import mimetypes
    
    filename = os.path.basename(doc_path)
    name_without_ext = os.path.splitext(filename)[0]
    docx_path = os.path.join(output_dir, f"{name_without_ext}.docx")
    
    # Méthode 1: Utiliser LibreOffice pour la conversion (préserve tableaux et mise en forme)
    try:
        # Chemin complet vers LibreOffice (plusieurs possibilités)
        libreoffice_paths = [
            'libreoffice',
            '/nix/store/i0x2skvhs1wbr5vffhhc53kd9jg2bp5q-libreoffice-7.6.4/bin/libreoffice', 
            '/usr/bin/libreoffice',
            '/usr/local/bin/libreoffice',
            '/opt/libreoffice/program/soffice'
        ]
        
        libreoffice_cmd = None
        for path in libreoffice_paths:
            try:
                subprocess.run([path, '--version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
                libreoffice_cmd = path
                break
            except (FileNotFoundError, subprocess.SubprocessError):
                continue
        
        if libreoffice_cmd:
            cmd = [
                libreoffice_cmd, '--headless', '--convert-to', 'docx', 
                '--outdir', output_dir, doc_path
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
            
            if os.path.exists(docx_path):
                print(f"Conversion réussie de {doc_path} en utilisant LibreOffice")
                return docx_path
        else:
            print("LibreOffice non trouvé dans les chemins standards")
            
    except (subprocess.SubprocessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
        print(f"Échec de la conversion via LibreOffice: {str(e)}")
    
    # Méthode 2: Essayer d'ouvrir directement avec python-docx
    # Cela fonctionne pour certains .docx renommés en .doc
    try:
        # Vérifier le type MIME pour s'assurer que c'est un document Word
        mime_type, _ = mimetypes.guess_type(doc_path)
        if mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            doc = Document(doc_path)
            doc.save(docx_path)
            print(f"Conversion directe réussie pour {doc_path}")
            return docx_path
        else:
            raise ValueError(f"file '{doc_path}' is not a Word file, content type is '{mime_type}'")
            
    except Exception as e:
        print(f"Échec de la conversion directe de {doc_path}: {str(e)}")
    
    # Méthode 3: Utiliser antiword pour extraire le texte (pour les vieux .doc)
    try:
        # Créer un fichier texte temporaire
        temp_txt = os.path.join(output_dir, f"{name_without_ext}.txt")
        
        cmd = ['antiword', doc_path]
        with open(temp_txt, 'w') as f:
            subprocess.run(cmd, stdout=f, stderr=subprocess.PIPE, check=True, timeout=30)
        
        # Créer un nouveau document DOCX à partir du texte extrait
        if os.path.exists(temp_txt) and os.path.getsize(temp_txt) > 0:
            doc = Document()
            
            # Ajouter le titre avec le nom du fichier
            doc.add_heading(f"Document: {filename}", level=1)
            
            # Lire le contenu du fichier texte et tenter de préserver les tableaux
            table_rows = []
            in_table = False
            
            with open(temp_txt, 'r') as f:
                for line in f:
                    line = line.strip()
                    
                    # Détection heuristique de tableau (lignes avec plusieurs | ou +)
                    if line and ('|' in line or '+' in line) and (line.count('|') > 2 or line.count('+') > 2):
                        if not in_table:
                            in_table = True
                            table_rows = []
                        
                        # Supprimer les caractères de bordure et récupérer les cellules
                        cells = [cell.strip() for cell in line.replace('+', '|').split('|') if cell.strip()]
                        if cells:
                            table_rows.append(cells)
                    else:
                        # Si nous sortions d'un tableau, créer le tableau dans le document
                        if in_table and table_rows:
                            table = doc.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
                            table.style = 'Table Grid'
                            
                            for i, row_data in enumerate(table_rows):
                                for j, cell_text in enumerate(row_data):
                                    if j < table.columns.length:  # Éviter l'index out of range
                                        table.cell(i, j).text = cell_text
                            
                            in_table = False
                            table_rows = []
                        
                        # Ajouter la ligne comme paragraphe si elle n'est pas vide
                        if line:
                            doc.add_paragraph(line)
            
            # Vérifier s'il reste un tableau à ajouter
            if in_table and table_rows:
                table = doc.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
                table.style = 'Table Grid'
                
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < table.columns.length:
                            table.cell(i, j).text = cell_text
            
            # Sauvegarder le document
            doc.save(docx_path)
            
            # Supprimer le fichier texte temporaire
            os.remove(temp_txt)
            
            print(f"Conversion réussie de {doc_path} en utilisant antiword avec extraction de tableaux")
            return docx_path
    
    except (subprocess.SubprocessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
        print(f"Échec de la conversion via antiword: {str(e)}")
        
    # Méthode 3b: Utiliser catdoc comme alternative à antiword
    try:
        # Créer un fichier texte temporaire
        temp_txt = os.path.join(output_dir, f"{name_without_ext}_catdoc.txt")
        
        cmd = ['catdoc', doc_path]
        with open(temp_txt, 'w') as f:
            subprocess.run(cmd, stdout=f, stderr=subprocess.PIPE, check=True, timeout=30)
        
        # Créer un nouveau document DOCX à partir du texte extrait - tentative de préservation des tableaux
        if os.path.exists(temp_txt) and os.path.getsize(temp_txt) > 0:
            doc = Document()
            
            # Ajouter le titre avec le nom du fichier
            doc.add_heading(f"Document: {filename}", level=1)
            
            # Variables pour la détection des tableaux
            table_rows = []
            in_table = False
            
            with open(temp_txt, 'r') as f:
                for line in f:
                    line = line.strip()
                    
                    # Détection heuristique de tableau (lignes structurées avec espaces réguliers)
                    if line and ('  ' in line or '\t' in line):
                        # Heuristique simple: si la ligne contient plusieurs espaces consécutifs
                        # et a une structure similaire aux lignes adjacentes, c'est probablement un tableau
                        if line.count('  ') >= 3 or line.count('\t') >= 3:
                            if not in_table:
                                in_table = True
                                table_rows = []
                            
                            # Diviser la ligne sur deux espaces ou plus pour obtenir les cellules
                            import re
                            cells = [cell.strip() for cell in re.split(r'  +|\t+', line) if cell.strip()]
                            if cells:
                                table_rows.append(cells)
                            continue
                    
                    # Si nous n'avons pas continué, nous ne sommes plus dans un tableau
                    if in_table:
                        if table_rows:
                            # Créer un tableau avec les lignes collectées
                            table = doc.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
                            table.style = 'Table Grid'
                            
                            for i, row_data in enumerate(table_rows):
                                for j, cell_text in enumerate(row_data):
                                    if j < table.columns.length:
                                        table.cell(i, j).text = cell_text
                        
                        in_table = False
                        table_rows = []
                    
                    # Ajouter la ligne comme paragraphe si elle n'est pas vide
                    if line:
                        doc.add_paragraph(line)
            
            # Vérifier s'il reste un tableau à ajouter
            if in_table and table_rows:
                table = doc.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
                table.style = 'Table Grid'
                
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < table.columns.length:
                            table.cell(i, j).text = cell_text
            
            # Sauvegarder le document
            doc.save(docx_path)
            
            # Supprimer le fichier texte temporaire
            os.remove(temp_txt)
            
            print(f"Conversion réussie de {doc_path} en utilisant catdoc avec extraction de tableaux")
            return docx_path
    
    except (subprocess.SubprocessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
        print(f"Échec de la conversion via catdoc: {str(e)}")
    
    # Méthode 4: En dernier recours, créer un document de substitution
    try:
        doc = Document()
        doc.add_heading(f"Document: {filename}", level=1)
        doc.add_paragraph("Ce document n'a pas pu être converti correctement.")
        doc.add_paragraph(f"Nom du fichier original: {filename}")
        doc.add_paragraph("Veuillez essayer d'ouvrir ce fichier directement avec votre traitement de texte.")
        doc.save(docx_path)
        
        print(f"Document de substitution créé pour {doc_path}")
        return docx_path
        
    except Exception as e:
        print(f"Échec de la création d'un document de substitution: {str(e)}")
        return None

def merge_docx_files(docx_files, output_path, status_dir):
    """
    Merge multiple .docx files into a single document
    
    Before each file's content, a header line with the filename is added.
    Updates status periodically.
    """
    try:
        # Créer un nouveau document pour la fusion
        master_doc = Document()
        
        # Initialisation des variables de status
        total_files = len(docx_files)
        processed = 0
        last_status_update = time.time()
        
        # Sauvegarder le statut initial
        save_status(status_dir, {
            "current_step": "merge",
            "complete": False,
            "file_count": total_files,
            "processed": processed,
            "total": total_files,
            "percent": 50
        })
        
        # Fusionner les documents
        for doc_path in docx_files:
            try:
                # Mettre à jour le compteur de traitement
                processed += 1
                
                # Mettre à jour périodiquement le statut (pas à chaque fichier pour améliorer les performances)
                current_time = time.time()
                if current_time - last_status_update > 1.0:  # Mise à jour toutes les secondes max
                    progress_percent = 50 + int((processed / total_files) * 30)  # 50-80% de la progression totale
                    save_status(status_dir, {
                        "current_step": "merge",
                        "complete": False,
                        "file_count": total_files,
                        "processed": processed,
                        "total": total_files,
                        "percent": progress_percent,
                        "status_text": f"Fusion du document {processed}/{total_files}..."
                    })
                    last_status_update = current_time
                
                # Extraire le nom du fichier
                filename = os.path.basename(doc_path)
                
                # Ajouter une ligne de séparation avec le nom du fichier
                separator_text = f"{filename}{'.' * 100}"
                separator = master_doc.add_paragraph(separator_text)
                separator.style = 'Heading 2'
                
                # Ouvrir le document à fusionner
                try:
                    doc = Document(doc_path)
                    
                    # Ajouter tous les paragraphes et tables du document à fusionner
                    for element in doc.element.body:
                        if element.tag.endswith('}p'):  # Paragraphe
                            new_p = master_doc.add_paragraph()
                            for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                                new_run = new_p.add_run()
                                for text in run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                    new_run.add_text(text.text)
                                # Préserver le formatage
                                if run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b') is not None:
                                    new_run.bold = True
                                if run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i') is not None:
                                    new_run.italic = True
                                if run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u') is not None:
                                    new_run.underline = True
                        
                        elif element.tag.endswith('}tbl'):  # Tableau
                            # Déterminer les dimensions du tableau
                            rows = element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
                            row_count = len(rows)
                            col_count = 0
                            for row in rows:
                                cells = row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                                col_count = max(col_count, len(cells))
                            
                            # Créer un nouveau tableau dans le document fusionné
                            if row_count > 0 and col_count > 0:
                                table = master_doc.add_table(rows=row_count, cols=col_count)
                                table.style = 'Table Grid'
                                
                                # Remplir le tableau avec les données
                                for i, row in enumerate(rows):
                                    cells = row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                                    for j, cell in enumerate(cells):
                                        if j < col_count:  # Éviter l'index out of range
                                            # Extraire le texte de la cellule
                                            cell_text = ''
                                            for text in cell.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                                cell_text += text.text
                                            # Affecter le texte à la cellule du nouveau tableau
                                            table.cell(i, j).text = cell_text
                
                except Exception as doc_error:
                    # En cas d'erreur dans un document spécifique, ajouter un message d'erreur
                    # et continuer avec les autres documents
                    error_paragraph = master_doc.add_paragraph()
                    error_run = error_paragraph.add_run(f"Erreur lors de la fusion du document {filename}: {str(doc_error)}")
                    error_run.bold = True
                    error_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Rouge
                    print(f"Erreur lors de la fusion de {doc_path}: {str(doc_error)}")
            
            except Exception as e:
                print(f"Erreur lors du traitement de {doc_path}: {str(e)}")
                # Continuer avec les autres documents
        
        # Mettre à jour le statut
        save_status(status_dir, {
            "current_step": "merging_complete",
            "complete": False,
            "file_count": total_files,
            "percent": 80,
            "status_text": "Fusion des fichiers terminée. Conversion en PDF..."
        })
        
        # Enregistrer le document fusionné
        master_doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        print(f"Erreur lors de la fusion des documents: {str(e)}")
        traceback.print_exc()
        
        save_status(status_dir, {
            "current_step": "error",
            "complete": False,
            "error": f"Erreur lors de la fusion des documents: {str(e)}",
            "status_text": f"Erreur lors de la fusion des documents: {str(e)}",
            "percent": 0
        })
        
        return None

def convert_docx_to_pdf(docx_path, pdf_path, status_dir):
    """
    Convert a .docx file to .pdf format
    
    This function attempts multiple methods to convert the document:
    1. libreoffice (if available)
    2. docx2pdf library (if installed)
    3. Basic fallback message if conversion is not possible
    """
    save_status(status_dir, {
        "current_step": "converting_to_pdf",
        "complete": False,
        "percent": 85,
        "status_text": "Conversion du document fusionné en PDF..."
    })
    
    # Méthode 1: Utiliser LibreOffice
    try:
        # Chemin complet vers LibreOffice (plusieurs possibilités)
        libreoffice_paths = [
            'libreoffice',
            '/nix/store/i0x2skvhs1wbr5vffhhc53kd9jg2bp5q-libreoffice-7.6.4/bin/libreoffice', 
            '/usr/bin/libreoffice',
            '/usr/local/bin/libreoffice',
            '/opt/libreoffice/program/soffice'
        ]
        
        libreoffice_cmd = None
        for path in libreoffice_paths:
            try:
                subprocess.run([path, '--version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
                libreoffice_cmd = path
                break
            except (FileNotFoundError, subprocess.SubprocessError):
                continue
                
        if libreoffice_cmd:
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                libreoffice_cmd, '--headless', '--convert-to', 'pdf', 
                '--outdir', output_dir, docx_path
            ]
            
            process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True, timeout=120)
            
            # Si le nom du fichier a changé, renommer le fichier PDF
            docx_basename = os.path.basename(docx_path)
            expected_pdf_name = os.path.splitext(docx_basename)[0] + '.pdf'
            expected_pdf_path = os.path.join(output_dir, expected_pdf_name)
            
            if os.path.exists(expected_pdf_path) and expected_pdf_path != pdf_path:
                os.rename(expected_pdf_path, pdf_path)
            
            if os.path.exists(pdf_path):
                save_status(status_dir, {
                    "current_step": "pdf_conversion_complete",
                    "complete": False,
                    "percent": 95,
                    "status_text": "Conversion PDF terminée. Finalisation..."
                })
                return pdf_path
        
    except (subprocess.SubprocessError, FileNotFoundError) as e:
        print(f"Échec de la conversion PDF via LibreOffice: {str(e)}")
    
    # Méthode 2: Essayer avec docx2pdf
    try:
        import docx2pdf
        docx2pdf.convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            save_status(status_dir, {
                "current_step": "pdf_conversion_complete",
                "complete": False,
                "percent": 95,
                "status_text": "Conversion PDF terminée. Finalisation..."
            })
            return pdf_path
            
    except ImportError:
        print("Bibliothèque docx2pdf non installée.")
    except Exception as e:
        print(f"Échec de la conversion PDF via docx2pdf: {str(e)}")
    
    # Méthode 3: Utiliser PyPDF2 et reportlab pour créer un PDF simple
    try:
        import io
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from PyPDF2 import PdfWriter
        
        # Lire le document DOCX et extraire le texte
        doc = Document(docx_path)
        text_content = []
        
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        # Créer un PDF avec reportlab
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        c.setFont("Helvetica", 12)
        
        y = 750  # Position verticale de départ
        line_height = 14
        
        for line in text_content:
            if y < 50:  # Si on est près du bas de la page
                c.showPage()
                c.setFont("Helvetica", 12)
                y = 750
            
            c.drawString(50, y, line)
            y -= line_height
        
        c.save()
        
        # Déplacer le pointeur au début
        packet.seek(0)
        
        # Créer un nouveau PDF avec PyPDF2
        writer = PdfWriter()
        writer.add_page(PdfWriter().append_pages_from_reader(packet).pages[0])
        
        # Écrire le fichier de sortie
        with open(pdf_path, 'wb') as outfile:
            writer.write(outfile)
        
        if os.path.exists(pdf_path):
            save_status(status_dir, {
                "current_step": "pdf_conversion_complete",
                "complete": False,
                "percent": 95,
                "status_text": "Conversion PDF terminée. Finalisation..."
            })
            return pdf_path
        
    except ImportError:
        print("Bibliothèques PyPDF2 ou reportlab non installées.")
    except Exception as e:
        print(f"Échec de la création d'un PDF via PyPDF2: {str(e)}")
    
    # Méthode 4: Créer un PDF basique avec reportlab seul
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica", 12)
        
        # Message d'information
        c.drawString(100, 700, "Conversion automatique du document DOCX en PDF")
        c.drawString(100, 680, "La conversion automatique n'a pas pu être effectuée.")
        c.drawString(100, 660, "Veuillez utiliser le fichier DOCX fourni.")
        
        c.save()
        
        if os.path.exists(pdf_path):
            save_status(status_dir, {
                "current_step": "pdf_conversion_complete",
                "complete": False,
                "percent": 95,
                "status_text": "Conversion PDF terminée avec un message d'information. Finalisation..."
            })
            return pdf_path
            
    except ImportError:
        print("Bibliothèque reportlab non installée.")
    except Exception as e:
        print(f"Échec de la création d'un PDF basique: {str(e)}")
    
    # Si tout échoue, créer un fichier texte simple
    try:
        txt_path = os.path.splitext(pdf_path)[0] + '.txt'
        with open(txt_path, 'w') as f:
            f.write("ERREUR DE CONVERSION PDF\n\n")
            f.write("La conversion automatique du document DOCX en PDF n'a pas pu être effectuée.\n")
            f.write("Veuillez utiliser le fichier DOCX fourni ou installer les outils nécessaires (LibreOffice, docx2pdf ou reportlab).")
        
        # Essayer une dernière fois avec reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            
            with open(txt_path, 'r') as f:
                y = 700
                for line in f:
                    c.drawString(100, y, line.strip())
                    y -= 20
            
            c.save()
            os.remove(txt_path)
            
        except Exception:
            # Si reportlab échoue, renommer le fichier texte en pdf
            if os.path.exists(txt_path):
                shutil.copy(txt_path, pdf_path)
        
        save_status(status_dir, {
            "current_step": "pdf_conversion_complete",
            "complete": False,
            "percent": 95,
            "status_text": "Conversion PDF échouée. Document texte créé à la place. Finalisation..."
        })
        
        return pdf_path
        
    except Exception as e:
        print(f"Échec complet de la conversion PDF: {str(e)}")
        
        save_status(status_dir, {
            "current_step": "error",
            "complete": False,
            "error": f"Erreur lors de la conversion PDF: {str(e)}",
            "status_text": f"Erreur lors de la conversion PDF: {str(e)}",
            "percent": 0
        })
        
        return None

def process_zip_file(zip_path, output_dir, status_dir=None, job_id=None):
    """
    Process a zip file containing .doc/.docx files:
    1. Extract all .doc and .docx files
    2. Convert .doc to .docx if needed
    3. Merge all into a single .docx
    4. Convert the merged file to PDF
    
    This function operates asynchronously and updates a status file.
    If job_id is provided, it will update the database with processing status.
    """
    # Créer une fonction qui sera exécutée dans un thread séparé
    def process_thread():
        try:
            start_time = time.time()
            
            # Créer les dossiers nécessaires
            job_dir = os.path.join(output_dir, f"{int(start_time)}_{job_id or ''}")
            os.makedirs(job_dir, exist_ok=True)
            
            extract_dir = os.path.join(job_dir, "extracted")
            os.makedirs(extract_dir, exist_ok=True)
            
            # Mise à jour du statut initial
            save_status(status_dir, {
                "current_step": "extract",
                "complete": False,
                "status_text": "Extraction des fichiers de l'archive ZIP...",
                "percent": 10
            })
            
            # Importer Flask pour avoir accès à la base de données
            if job_id:
                try:
                    from flask import current_app
                    with current_app.app_context():
                        # Importer le modèle de base de données
                        from models import ProcessingJob, db
                        
                        # Mettre à jour le statut du job dans la base de données
                        job = ProcessingJob.query.filter_by(job_id=job_id).first()
                        if job:
                            job.status = "processing"
                            db.session.commit()
                except ImportError:
                    print("Impossible d'importer Flask pour mettre à jour la base de données")
            
            # Étape 1: Extraire les fichiers .doc et .docx
            extracted_files = extract_doc_files(zip_path, extract_dir)
            file_count = len(extracted_files)
            
            if file_count == 0:
                error_msg = "Aucun fichier .doc ou .docx trouvé dans l'archive ZIP."
                save_status(status_dir, {
                    "current_step": "error",
                    "complete": False,
                    "error": error_msg,
                    "status_text": error_msg,
                    "percent": 0
                })
                
                # Mettre à jour le statut dans la base de données
                if job_id:
                    try:
                        from flask import current_app
                        with current_app.app_context():
                            # Importer le modèle de base de données
                            from models import ProcessingJob, db
                            
                            # Mettre à jour le statut du job dans la base de données
                            job = ProcessingJob.query.filter_by(job_id=job_id).first()
                            if job:
                                job.status = "error"
                                job.completed_at = datetime.utcnow()
                                db.session.commit()
                    except ImportError:
                        pass
                
                return None
            
            # Mise à jour du statut avec le nombre de fichiers
            save_status(status_dir, {
                "current_step": "convert",
                "complete": False,
                "file_count": file_count,
                "status_text": "Conversion des fichiers DOC en DOCX...",
                "percent": 30
            })
            
            # Étape 2: Convertir les .doc en .docx si nécessaire
            docx_files = []
            converted = 0
            
            for doc_path in extracted_files:
                # Mise à jour périodique du statut
                if file_count > 10 and converted % 10 == 0:
                    progress_percent = 30 + int((converted / file_count) * 20)  # 30-50% de la progression totale
                    save_status(status_dir, {
                        "current_step": "convert",
                        "complete": False,
                        "file_count": file_count,
                        "converted": converted,
                        "total_files": file_count,
                        "progress_percent": progress_percent,
                        "status_text": f"Conversion des fichiers ({converted}/{file_count})...",
                        "percent": progress_percent
                    })
                
                if doc_path.lower().endswith('.doc'):
                    # Convertir le .doc en .docx
                    docx_path = convert_doc_to_docx(doc_path, extract_dir)
                    if docx_path:
                        docx_files.append(docx_path)
                elif doc_path.lower().endswith('.docx'):
                    # Déjà au format .docx
                    docx_files.append(doc_path)
                
                converted += 1
            
            # Mise à jour du statut pour la fusion
            save_status(status_dir, {
                "current_step": "merge",
                "complete": False,
                "file_count": file_count,
                "status_text": "Fusion des documents en un seul fichier...",
                "percent": 50
            })
            
            # Vérifier qu'il y a des fichiers à fusionner
            if not docx_files:
                error_msg = "Aucun fichier n'a pu être converti correctement."
                save_status(status_dir, {
                    "current_step": "error",
                    "complete": False,
                    "error": error_msg,
                    "status_text": error_msg,
                    "percent": 0
                })
                
                # Mettre à jour le statut dans la base de données
                if job_id:
                    try:
                        from flask import current_app
                        with current_app.app_context():
                            from models import ProcessingJob, db
                            job = ProcessingJob.query.filter_by(job_id=job_id).first()
                            if job:
                                job.status = "error"
                                job.completed_at = datetime.utcnow()
                                db.session.commit()
                    except ImportError:
                        pass
                
                return None
            
            # Étape 3: Fusionner tous les fichiers .docx
            output_docx = os.path.join(job_dir, "merged.docx")
            merged_docx = merge_docx_files(docx_files, output_docx, status_dir)
            
            if not merged_docx or not os.path.exists(merged_docx):
                error_msg = "Échec de la fusion des documents."
                save_status(status_dir, {
                    "current_step": "error",
                    "complete": False,
                    "error": error_msg,
                    "status_text": error_msg,
                    "percent": 0
                })
                
                # Mettre à jour le statut dans la base de données
                if job_id:
                    try:
                        from flask import current_app
                        with current_app.app_context():
                            from models import ProcessingJob, db
                            job = ProcessingJob.query.filter_by(job_id=job_id).first()
                            if job:
                                job.status = "error"
                                job.completed_at = datetime.utcnow()
                                db.session.commit()
                    except ImportError:
                        pass
                
                return None
            
            # Étape 4: Convertir le fichier fusionné en PDF
            output_pdf = os.path.join(job_dir, "merged.pdf")
            pdf_path = convert_docx_to_pdf(merged_docx, output_pdf, status_dir)
            
            # Finaliser le traitement
            end_time = time.time()
            processing_time = int(end_time - start_time)
            
            # Mise à jour du statut final
            save_status(status_dir, {
                "current_step": "complete",
                "complete": True,
                "file_count": file_count,
                "output_docx": "merged.docx",
                "output_pdf": "merged.pdf" if pdf_path else None,
                "start_time": int(start_time),
                "end_time": int(end_time),
                "processing_time": processing_time,
                "stats": {
                    "processing_time": processing_time,
                    "file_count": file_count
                },
                "status_text": "Traitement terminé avec succès.",
                "percent": 100
            })
            
            # Mettre à jour le statut dans la base de données
            if job_id:
                try:
                    from flask import current_app
                    with current_app.app_context():
                        from models import ProcessingJob, UsageStat, db
                        
                        # Mettre à jour le job
                        job = ProcessingJob.query.filter_by(job_id=job_id).first()
                        if job:
                            job.status = "completed"
                            job.completed_at = datetime.utcnow()
                            job.file_count = file_count
                            job.processing_time = processing_time
                            db.session.commit()
                        
                        # Mettre à jour les statistiques d'utilisation
                        today = datetime.utcnow().date()
                        stat = UsageStat.query.filter_by(date=today).first()
                        
                        if stat:
                            stat.total_jobs += 1
                            stat.total_files_processed += file_count
                            stat.total_processing_time += processing_time
                        else:
                            new_stat = UsageStat(
                                date=today,
                                total_jobs=1,
                                total_files_processed=file_count,
                                total_processing_time=processing_time
                            )
                            db.session.add(new_stat)
                        
                        db.session.commit()
                except ImportError:
                    print("Impossible d'importer Flask pour mettre à jour la base de données")
                except Exception as db_error:
                    print(f"Erreur lors de la mise à jour de la base de données: {str(db_error)}")
            
            return {
                "job_dir": job_dir,
                "docx_path": merged_docx,
                "pdf_path": pdf_path,
                "file_count": file_count,
                "processing_time": processing_time
            }
            
        except Exception as e:
            error_msg = f"Erreur lors du traitement de l'archive ZIP: {str(e)}"
            traceback.print_exc()
            
            save_status(status_dir, {
                "current_step": "error",
                "complete": False,
                "error": error_msg,
                "status_text": error_msg,
                "percent": 0
            })
            
            # Mettre à jour le statut dans la base de données
            if job_id:
                try:
                    from flask import current_app
                    with current_app.app_context():
                        from models import ProcessingJob, db
                        job = ProcessingJob.query.filter_by(job_id=job_id).first()
                        if job:
                            job.status = "error"
                            job.completed_at = datetime.utcnow()
                            db.session.commit()
                except ImportError:
                    pass
                except Exception as db_error:
                    print(f"Erreur lors de la mise à jour de la base de données: {str(db_error)}")
            
            return None
    
    # Démarrer le traitement dans un thread séparé
    thread = threading.Thread(target=process_thread)
    thread.daemon = True
    thread.start()
    
    return True

def cleanup_old_files(directory, max_age_hours=24):
    """Delete files older than max_age_hours from the directory"""
    try:
        cutoff_time = datetime.now() - timedelta(hours=max_age_hours)
        
        for root, dirs, files in os.walk(directory):
            for dir_name in dirs:
                dir_path = os.path.join(root, dir_name)
                try:
                    # Vérifier si le nom du dossier commence par un timestamp
                    if dir_name.split('_')[0].isdigit():
                        # Convertir le timestamp en datetime
                        timestamp = int(dir_name.split('_')[0])
                        dir_time = datetime.fromtimestamp(timestamp)
                        
                        # Supprimer si plus ancien que la limite
                        if dir_time < cutoff_time:
                            shutil.rmtree(dir_path, ignore_errors=True)
                            print(f"Suppression du dossier ancien: {dir_path}")
                except (ValueError, IndexError, OSError) as e:
                    print(f"Erreur lors du nettoyage du dossier {dir_name}: {str(e)}")
                    continue
        
        return True
    except Exception as e:
        print(f"Erreur lors du nettoyage des fichiers anciens: {str(e)}")
        return False